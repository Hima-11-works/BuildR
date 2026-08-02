# ──────────────────────────────────────────────────────────────
# services/parser_service.py — Text extraction from PDF/DOCX
# ──────────────────────────────────────────────────────────────

from __future__ import annotations

import io
import re

# ── Heavy imports: deferred ──────────────────────────────────
# `pypdf` pulls in ~7 MB at import time (PDF parsing tables,
# glyph lists, cryptography fallbacks). `python-docx` adds more.
# Both are only used in two narrow functions, so we defer the
# imports to first call. Routes that never parse a resume (e.g.
# auth, profile, tailoring, library) avoid the cost entirely.


def extract_text_from_pdf(stream: io.BytesIO) -> str:
    """
    Extract raw text from a PDF file stream using pypdf.
    """
    # ── INSTRUMENTATION ──
    # Logs the number of PDF pages after the parser has them loaded
    # but before they are walked for text. Tells us whether a
    # 200-page PDF (vs a normal 1-2 page resume) is reaching the
    # worker.
    from services._diagnostics import log_event
    try:
        # Lazy import: pypdf + its deps (~7 MB) load only when the
        # first resume-PDF upload is parsed.
        import pypdf
        reader = pypdf.PdfReader(stream)
        try:
            num_pages = len(reader.pages)
        except Exception:  # noqa: BLE001 — pypdf can fail to count on encrypted PDFs
            num_pages = "?"
        log_event("pdf_pages_loaded", num_pages=num_pages)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {str(e)}")


def extract_text_from_docx(stream: io.BytesIO) -> str:
    """
    Extract raw text from a DOCX file stream using python-docx.
    Includes text from paragraphs and tables.
    """
    from services._diagnostics import log_event
    try:
        # Lazy import: python-docx + its deps load only when the
        # first resume-DOCX upload is parsed.
        import docx
        d = docx.Document(stream)
        try:
            num_paragraphs = len(d.paragraphs)
            num_tables = len(d.tables)
        except Exception:  # noqa: BLE001
            num_paragraphs = "?"
            num_tables = "?"
        log_event("docx_structure_loaded",
                  num_paragraphs=num_paragraphs,
                  num_tables=num_tables)
        text_parts = []

        # Extract from paragraphs
        for paragraph in d.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        # Extract from tables
        for table in d.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells if cell.text]
                if row_cells:
                    text_parts.append(" ".join(row_cells))

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


# ── List-to-HTML conversion ───────────────────────────────────
# Patterns recognised as list item prefixes.
# Group 1 (optional) captures a numeric prefix such as "1." or "1)".
# Group 2 captures the item text that follows the marker.

_BULLET_RE = re.compile(
    r"^[ \t]*"                       # optional leading whitespace
    r"(?:"
        r"(?P<num>\d+[.)]\s+)"       # numbered: "1. " or "1) "
        r"|"
        r"(?:[•·▪▸►\-\*–—]\s+)"     # bullet chars: •, -, *, –, —, ▪ …
    r")"
    r"(?P<text>.+)$",
    re.MULTILINE,
)

_NUMBERED_LINE_RE = re.compile(r"^\d+[.)]\s+", re.MULTILINE)


def _escape_html(text: str) -> str:
    """Minimal HTML-escape for plain text being wrapped in tags."""
    return (
        text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def text_to_html(text: str) -> str:
    """
    Convert a plain-text string that may contain list patterns into semantic
    HTML suitable for the app's rich-text editor fields.

    Conversion rules (applied in order):
    1. Already-HTML — if the string starts with a block-level tag (<ul>, <ol>,
       <li>, <p>, <br>) it is returned unchanged; the editor content is intact.
    2. Bullet lists — if the majority of non-empty lines begin with a common
       bullet marker (•, -, *, –, —, ▪, ►, ▸) the block is wrapped in <ul>.
    3. Numbered lists — if the majority of non-empty lines begin with "N." or
       "N)" the block is wrapped in <ol>.
    4. Multi-line plain text — if there are multiple non-empty lines but no
       clear list marker, line breaks are preserved using <br> tags so the
       content is not collapsed into a single paragraph.
    5. Single-line text — returned as-is (no wrapping needed).

    Items extracted from lists have their leading marker stripped before
    being placed in <li> tags.
    """
    if not text or not text.strip():
        return text

    # Rule 1 — already HTML
    stripped = text.strip()
    if re.match(r"^<(ul|ol|li|p|br|div|h[1-6])", stripped, re.IGNORECASE):
        return text

    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return text

    if len(lines) == 1:
        # Single logical line — just return clean text
        return stripped

    # Check what fraction of lines look like list items.
    bullet_lines = [
        ln for ln in lines
        if re.match(r"^[ \t]*[•·▪▸►\-\*–—]\s+", ln)
    ]
    numbered_lines = [
        ln for ln in lines
        if re.match(r"^[ \t]*\d+[.)]\s+", ln)
    ]

    bullet_ratio  = len(bullet_lines)  / len(lines)
    numbered_ratio = len(numbered_lines) / len(lines)

    # Require ≥ 50 % of lines to carry a marker before treating as a list.
    # This avoids mis-classifying paragraphs that happen to start with a dash.

    if bullet_ratio >= 0.5:
        # Rule 2 — unordered list
        items = []
        for ln in lines:
            content = re.sub(r"^[ \t]*[•·▪▸►\-\*–—]\s+", "", ln).strip()
            if content:
                items.append(f"<li>{_escape_html(content)}</li>")
        if items:
            return "<ul>" + "".join(items) + "</ul>"

    if numbered_ratio >= 0.5:
        # Rule 3 — ordered list
        items = []
        for ln in lines:
            content = re.sub(r"^[ \t]*\d+[.)]\s+", "", ln).strip()
            if content:
                items.append(f"<li>{_escape_html(content)}</li>")
        if items:
            return "<ol>" + "".join(items) + "</ol>"

    # Rule 4 — multi-line plain text: preserve breaks
    return "<br>".join(_escape_html(ln.strip()) for ln in lines if ln.strip())


def postprocess_parsed_profile(profile: "Profile") -> "Profile":  # type: ignore[name-defined]
    """
    Apply rich-text list formatting to every free-text field in a Profile that
    was returned by the AI parser.

    Fields processed
    ----------------
    - profile.achievements          (str — rendered by the HTML editor)
    - profile.certifications[*].description  (Optional[str])

    The experience bullets and project bullets are already structured as
    list[str] by the schema, so they do not need conversion here.

    Returns the same Profile object (mutated in-place) for convenience.
    """
    from models.profile import Certification

    # ── Achievements ─────────────────────────────────────────────
    if profile.achievements and isinstance(profile.achievements, str):
        profile.achievements = text_to_html(profile.achievements)

    # ── Certification descriptions ────────────────────────────────
    updated_certs = []
    for cert in profile.certifications:
        if cert.description and isinstance(cert.description, str):
            new_desc = text_to_html(cert.description)
            if new_desc != cert.description:
                cert = Certification(
                    name=cert.name,
                    issuer=cert.issuer,
                    date=cert.date,
                    description=new_desc,
                )
        updated_certs.append(cert)
    profile.certifications = updated_certs

    return profile
